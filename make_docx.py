import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\Users\berge\OneDrive\Skrivebord\Claude code\github-pages-deploy"
MD   = os.path.join(BASE, "dikult216_final.md")
OUT  = os.path.join(BASE, "dikult216_final.docx")

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# --- Default font ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    p.add_run("Oppdater feltet i Word for å vise innholdsfortegnelsen.")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p.add_run()._r.append(fld_end)

def set_heading(p, level):
    sizes = {1: 18, 2: 14, 3: 12}
    p.style = doc.styles[f"Heading {level}"]
    for run in p.runs:
        run.font.size = Pt(sizes.get(level, 12))

def add_inline(paragraph, text):
    """Parse **bold** and *italic* in text and add runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

with open(MD, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip("\n")

    # HTML comment helpers for export formatting
    if line.strip() == "<!-- PAGEBREAK -->":
        doc.add_page_break()
        i += 1
        continue

    if line.strip() == "<!-- TOC -->":
        add_toc(doc)
        i += 1
        continue

    center_match = re.match(r'^<div align="center">(.*)</div>$', line.strip())
    if center_match:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, center_match.group(1))
        i += 1
        continue

    # H1
    if line.startswith("# ") and not line.startswith("## "):
        p = doc.add_heading(line[2:], level=1)
        i += 1
        continue

    # H2
    if line.startswith("## ") and not line.startswith("### "):
        p = doc.add_heading(line[3:], level=2)
        i += 1
        continue

    # H3
    if line.startswith("### "):
        p = doc.add_heading(line[4:], level=3)
        i += 1
        continue

    # HR
    if line.strip() == "---":
        doc.add_paragraph()
        i += 1
        continue

    # Image  ![alt](filename)  optionally followed by *caption*
    img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if img_match:
        img_file = os.path.join(BASE, img_match.group(2))
        if os.path.exists(img_file):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_file, width=Inches(4.5))
        else:
            doc.add_paragraph(f"[Bilde ikke funnet: {img_match.group(2)}]")
        i += 1
        # Check if next non-empty line is a caption (*italic*)
        while i < len(lines) and lines[i].strip() == "":
            i += 1
        if i < len(lines):
            cap_line = lines[i].strip()
            cap_match = re.match(r"^\*(.+)\*$", cap_line)
            if cap_match:
                add_caption(doc, cap_match.group(1))
                i += 1
        continue

    # Blank line
    if line.strip() == "":
        i += 1
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    # Strip leading *text* that is a standalone italic block (standalone captions already handled above)
    add_inline(p, line)
    i += 1

doc.save(OUT)
print(f"Saved: {OUT}")
